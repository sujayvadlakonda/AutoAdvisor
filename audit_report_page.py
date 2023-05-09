import tkinter as tk
import os
import docx
from tkinter import ttk
from tkinter.filedialog import asksaveasfilename
from tkinter import messagebox as mbox
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from transcript import Transcript
from audit_report import AuditReport
from track import DataScience, CyberSecurity, IntelligentSystems, InteractiveComputing, Network, Systems
from track import Traditional, SoftwareEngineering
from course import LevelingCourse, LevelingCourseStatus
from course_finder import get_courses


class AuditReportPage(ttk.Frame):
    def __init__(self, container, controller):
        super().__init__(container)
        self.controller = controller
        self.transcript_path = tk.StringVar()  # holds path to transcript

        self.student_name = ""
        self.student_id = ""
        self.student_plan = ""
        self.student_major = ""
        self.student_track = tk.StringVar()

        self.core_gpa = ""
        self.elective_gpa = ""
        self.overall_gpa = ""

        self.core_courses = ""
        self.elective_courses = ""

        self.disposition_dict = {
            "dp_pre_req_class": [],  # holds classes to display in gui
            "disp_options": ["None", "Completed", "Waived", "Not required by plan or elective", "Other"],
            "disp_selections": [],  # holds user selected disposition choice
            "opt_menu": [],  # holds disposition menu widget
            "user_course_comment": [],  # Hold's user's entry in text entry box
            "entry_box": [],  # holds entry text box widget
            "lbl_course_widget": [],  # holds course name label widget
            "lbl_comment_widget": [],  # holds comment label widget
            "btn_submit": []  # holds submit button widget
        }

        # Holds strings to be printed for the Outstanding Requirements section of the Word doc
        self.out_req_sect = ""

        self.select_elective = tk.IntVar()  # holds selection value to Student Taking Extra Elective question
        self.xtra_elective_result = ""  # saves selection value to Student Taking Extra Elective question

        self.select_addi_class = tk.IntVar()  # holds selection value to Student Taking Additional Courses question
        self.addi_course_question_result = ""  # saves selection value to Student Taking Additional Courses question
        self.class_quantity = tk.StringVar()  # Holds quantity value to Quantity of Additional Courses question
        self.class_quantity_output = ""  # Saves quality value ot Quantity of Additional Courses question

        # Creates Instances of gpa related gui widgets
        self.lbl_core_gpa_display = ttk.Label()
        self.lbl_elective_gpa_display = ttk.Label()
        self.lbl_overall_gpa_display = ttk.Label()

        # Instances attribute of Quantity Additional Courses entry related widget
        self.lbl_quan_class = ttk.Label()
        self.class_quantity_box = ttk.Entry()
        self.lbl_enter = ttk.Label()

        self.lbl_out_req_sect = ttk.Label()  # outstanding requirements widget instance attribute

        # Instance attribute of GUI buttons
        self.previous_btn = ttk.Button()
        self.audit_report_btn = ttk.Button()
        self.homepage_btn = ttk.Button()

        # Handles gui for audit report page
        self.style = ttk.Style(self)
        self.style.configure("aud_report_gui.TFrame", background="white")
        self.style.configure("BlWht.TLabel", font=("Segoe UI", 20), foreground="black", background="white")
        self.style.configure("BlackSmall.TLabel", font=("Arial", 14), foreground="black", background="orange")
        self.style.configure("instruct_txt.TLabel", font=("Georgia", 14), foreground="black", background="orange")
        self.style.configure("Black_txt.TLabel", font=("Calibri", 14), foreground="black", background="yellow")
        self.style.configure("gray_subtext.TLabel", font=("Calibri", 12), foreground="#4C4E52", background="white")
        self.style.configure("black_subtext.TLabel", font=("Calibri", 12), foreground="black", background="white")
        self.style.configure("input.TRadiobutton", background="pink")

        # Handles frame expansion when application window is expanded
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)

        # Frame background design
        frame = ttk.Frame(self, style="aud_report_gui.TFrame")

        # Background Frame's page space distribution starting from row 0, column 0
        for row_index in range(49):
            frame.grid_rowconfigure(row_index, weight=1)
        for col_index in range(15):
            frame.grid_columnconfigure(col_index, weight=1)
        frame.grid(column=0, row=0, sticky="nsew", columnspan=16)

        # Canvas to hold scrollbar
        canvas = tk.Canvas(frame, bg="white", bd=0, highlightthickness=0)
        canvas.grid(column=0, row=0, sticky="nsew", columnspan=15, rowspan=50)  # positioning

        # Page Scrollbar and design
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollbar.grid(row=0, column=15, rowspan=50, sticky="ns")  # positioning

        # Frame to hold contents to scroll
        self.scrollable_frame = ttk.Frame(canvas, style="aud_report_gui.TFrame")
        self.scrollable_frame["padding"] = (5, 0, 5, 0)  # adds padding for spacing aesthetic

        # Scrollable frame page space distribution
        for row_index in range(49):
            self.scrollable_frame.grid_rowconfigure(row_index, weight=1)
        for col_index in range(15):
            self.scrollable_frame.grid_columnconfigure(col_index, weight=1)

        # Canvas Design and expansion when window expands
        canvas_win = canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")  # positioning
        canvas.bind('<Configure>', lambda e: canvas.itemconfig(canvas_win, width=e.width))
        self.scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.configure(yscrollcommand=scrollbar.set)  # sets scrollbar

        # Handles scrolling the frame based on OS when hovering over the scrollbar
        if os.name == "posix":
            self.scrollable_frame.bind(
                "<Enter>",
                lambda e: self.scrollable_frame.bind_all('<MouseWheel>',
                                                         canvas.yview_scroll(int(e.delta), "units")))
            self.scrollable_frame.bind("<Leave>", lambda e: self.scrollable_frame.unbind_all('<MouseWheel>'))
        else:
            self.scrollable_frame.bind(
                "<Enter>",
                lambda e: self.scrollable_frame.bind_all('<MouseWheel>',
                                                         canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")))
            self.scrollable_frame.bind("<Leave>", lambda e: self.scrollable_frame.unbind_all('<MouseWheel>'))

        # Start Audit Report Process Button
        btn_start_audit = ttk.Button(
            self.scrollable_frame,
            text="Start Audit",
            command=self.audit_report_gui  # Call function to display contents of gui frame
        )
        btn_start_audit.grid(column=0, row=0, columnspan=1, pady=(5, 0))  # positioning

    # Saves value of Taking Extra Elective Question selection
    def xtra_elective_question_value(self):
        self.xtra_elective_result = self.select_elective.get()  # saves "value" of user selection

    # Asks user for Additional Graduate Courses quantity if user selected yes to Additional Graduate Courses question
    def addi_course_question_value(self):
        self.addi_course_question_result = self.select_addi_class.get()  # saves value of user selection
        class_amount_question = "If yes, enter additional graduate course quantity:"
        submit_quantity_entry = "[Press Enter to Submit]"

        # Prompts user for additional course quantity
        if self.select_addi_class.get() == 1:
            self.lbl_quan_class = ttk.Label(self.scrollable_frame, text=class_amount_question, style="Black_txt.TLabel")
            self.lbl_quan_class.grid(column=0, row=14, columnspan=3, sticky="w", pady=(5, 10))  # text positioning
            # Entry text box for user to enter how many additional graduate courses the student is taking
            self.class_quantity_box.grid(column=3, row=14, columnspan=1, sticky="w", pady=(5, 10))  # positioning
            self.class_quantity_box.bind("<Return>", self.addi_course_entry)  # saves entry upon user press Enter
            # Displays instructions for user to press enter to save their entry
            self.lbl_enter = ttk.Label(self.scrollable_frame, text=submit_quantity_entry, style="gray_subtext.TLabel")
            self.lbl_enter.grid(column=4, row=14, columnspan=2, sticky="w", pady=(5, 10))  # positioning

    # Saves value entered in Additional Graduate Course text entry box upon the Enter key being pressed
    def addi_course_entry(self, event):
        self.class_quantity_output = self.class_quantity.get()

    # Ensures updated user selected disposition values are tracked upon selection in the GUI
    def disposition_update(self, *args):
        disp_tracker = []
        for saved_disp in self.disposition_dict["disp_selections"]:
            disp_status_word = saved_disp.get()
            disp_tracker.append(disp_status_word)

    # Ensures text entry box values are tracked upon clicking the submit button
    def comment_tracker(self):
        comment_tracker = []
        for entry in self.disposition_dict["user_course_comment"]:
            comment_word = entry.get()
            comment_tracker.append(comment_word)

    # Opens File Explorer to save file
    def save_file(self):
        default_file_name = "audit_report"
        word_doc = "Microsoft Word Document"
        file_types = [("All Files", "*.*"), (word_doc, "*.docx*"), ]
        save_file_progress = "Download Complete. Please check the audit report before exiting the program."

        # Opens File Explorer based on OS
        if os.name == 'posix':
            audit_report_filepath = asksaveasfilename(
                initialdir="/",
                title="Save As",
                defaultextension=".docx",
                initialfile=default_file_name)
            # Checks if user pressed Save or Cancel button
            if not audit_report_filepath:
                mbox.showerror("Error", "File Not Saved. Please, try Again.")  # file not saved error message
            else:
                self.generate_report_content(audit_report_filepath)  # calls function to write text to file
                # Displays message to user that the file is being saved
                mbox.showinfo("Information", save_file_progress)
        else:
            audit_report_filepath = asksaveasfilename(
                defaultextension=".docx",
                filetypes=file_types,
                initialdir="\\",
                initialfile=default_file_name,
                title="Save As")
            # Checks if user pressed cancel or save in file explorer
            if not audit_report_filepath:
                mbox.showerror("Error", "File Not Saved. Please, try Again.")  # file not saved error message
            else:
                self.generate_report_content(audit_report_filepath)  # calls function to write text to file
                # Displays message to user that the file is being saved
                mbox.showinfo("Information", save_file_progress)

    # Displays Contents of Audit Report Page GUI
    def audit_report_gui(self):
        # resets values if already present
        if self.core_gpa:
            self.lbl_core_gpa_display.grid_forget()
            self.lbl_elective_gpa_display.grid_forget()
            self.lbl_overall_gpa_display.grid_forget()
            self.lbl_out_req_sect.grid_forget()
            for widget in self.disposition_dict["lbl_course_widget"]:
                widget.grid_forget()
            for widget in self.disposition_dict["opt_menu"]:
                widget.grid_forget()
            for widget in self.disposition_dict["lbl_comment_widget"]:
                widget.grid_forget()
            for widget in self.disposition_dict["entry_box"]:
                widget.grid_forget()
            for widget in self.disposition_dict["btn_submit"]:
                widget.grid_forget()
            self.previous_btn.grid_forget()
            self.audit_report_btn.grid_forget()
            self.homepage_btn.grid_forget()
        if self.select_elective.get() == 1 or 2:
            self.select_elective.set(0)
        if self.select_addi_class.get() == 1 or 2:
            if self.select_addi_class.get() == 1:
                self.lbl_quan_class.grid_forget()
                self.class_quantity_box.grid_forget()
                self.lbl_enter.grid_forget()
            self.select_addi_class.set(0)
        if self.class_quantity.get():
            self.class_quantity_box.delete(0, "end")
        self.disposition_dict["dp_pre_req_class"].clear()
        self.disposition_dict["disp_selections"].clear()
        self.disposition_dict["opt_menu"].clear()
        self.disposition_dict["user_course_comment"].clear()
        self.disposition_dict["entry_box"].clear()
        self.disposition_dict["lbl_course_widget"].clear()
        self.disposition_dict["lbl_comment_widget"].clear()
        self.disposition_dict["btn_submit"].clear()

        transcript_filepath = self.transcript_path  # holds filepath of uploaded file
        track_class = self.check_track()  # holds class for corresponding selected track
        audit_report = AuditReport(transcript_filepath, track_class)  # used to call audit report related methods
        self.core_gpa = audit_report.get_core_gpa_section()  # holds core gpa displayed in audit report
        self.elective_gpa = audit_report.get_elective_gpa_section()  # holds elective gpa displayed in elective gpa
        self.overall_gpa = audit_report.get_combined_gpa_section()  # holds overall gpa displayed in elective gpa
        self.out_req_sect = audit_report.get_outstanding_requirements_section()  # holds outstanding requirements text

        # Audit Report page title and design
        lbl_aud_report = ttk.Label(self.scrollable_frame, text="Audit Report", style="BlWht.TLabel")
        lbl_aud_report.grid(column=1, row=0, columnspan=2, sticky="nw", pady=(5, 0))  # text positioning

        # Displays core, elective, and overall gpa label on screen
        lbl_core_gpa = ttk.Label(self.scrollable_frame, text="Core GPA: ", style="BlackSmall.TLabel")
        lbl_core_gpa.grid(column=0, row=6, columnspan=1, sticky="w", pady=(20, 0))  # text positioning
        lbl_elective_gpa = ttk.Label(self.scrollable_frame, text="Elective GPA: ", style="BlackSmall.TLabel")
        lbl_elective_gpa.grid(column=0, row=7, columnspan=1, sticky="w")  # text positioning
        lbl_overall_gpa = ttk.Label(self.scrollable_frame, text="Overall GPA: ", style="BlackSmall.TLabel")
        lbl_overall_gpa.grid(column=0, row=8, columnspan=1, sticky="w", pady=(0, 20))  # text positioning

        # Displays core, elective, and overall gpa on screen
        self.lbl_core_gpa_display = ttk.Label(self.scrollable_frame, text=self.core_gpa, style="Black_txt.TLabel")
        self.lbl_elective_gpa_display = ttk.Label(self.scrollable_frame, text="", style="Black_txt.TLabel")
        self.lbl_elective_gpa_display.configure(text=self.elective_gpa)  # sets text
        self.lbl_overall_gpa_display = ttk.Label(self.scrollable_frame, text=self.overall_gpa, style="Black_txt.TLabel")
        self.lbl_core_gpa_display.grid(column=1, row=6, columnspan=1, sticky="w", pady=(20, 0))  # text positioning
        self.lbl_elective_gpa_display.grid(column=1, row=7, columnspan=1, sticky="w")  # text positioning
        self.lbl_overall_gpa_display.grid(column=1, row=8, columnspan=1, sticky="w", pady=(0, 20))  # text positioning

        # Prompts user to answer questions label and design
        prompt_user_instruct = "Answer the Following Questions:"
        lbl_xtra_elect_prompt = ttk.Label(self.scrollable_frame, text=prompt_user_instruct, style="instruct_txt.TLabel")
        lbl_xtra_elect_prompt.grid(column=0, row=11, columnspan=2, sticky="w", pady=(20, 0))  # positioning

        # Asks user if student will be taking extra electives (used for elective gpa calculations)
        elective_question = "Student taking extra electives:"
        lbl_xtra_elect = ttk.Label(self.scrollable_frame, text=elective_question, style="BlackSmall.TLabel")
        lbl_xtra_elect.grid(column=0, row=12, columnspan=2, sticky="w", pady=(10, 0))  # positioning

        # Taking Extra Electives question (yes/no) selection options
        elective_confirm = ttk.Radiobutton(
            self.scrollable_frame,
            text="Yes",
            value=1,
            variable=self.select_elective,
            command=self.xtra_elective_question_value,
            style="input.TRadiobutton"
        )
        elective_confirm.grid(column=2, row=12, columnspan=1, sticky="w", pady=(10, 0))  # positioning
        elective_deny = ttk.Radiobutton(
            self.scrollable_frame,
            text="No",
            value=2,
            variable=self.select_elective,
            command=self.xtra_elective_question_value,
            style="input.TRadiobutton"
        )
        elective_deny.grid(column=3, row=12, columnspan=1, sticky="w", pady=(10, 0))  # positioning

        # Asks user if student is taking additional graduate courses and how many (used for gpa calculations)
        grad_course_question = "Student taking additional graduate courses:"
        lbl_grad_class = ttk.Label(self.scrollable_frame, text=grad_course_question, style="BlackSmall.TLabel")
        lbl_grad_class.grid(column=0, row=13, columnspan=3, sticky="w", pady=(20, 0))  # text positioning

        # Taking Additional Graduate Courses (yes/no) question selection options
        addi_courses_confirm = ttk.Radiobutton(
            self.scrollable_frame,
            text="Yes",
            value=1,
            variable=self.select_addi_class,
            command=self.addi_course_question_value,
            style="input.TRadiobutton"
        )
        addi_courses_confirm.grid(column=3, row=13, columnspan=1, sticky="w", pady=(20, 0))  # positioning
        addi_courses_deny = ttk.Radiobutton(
            self.scrollable_frame,
            text="No",
            value=2,
            variable=self.select_addi_class,
            command=self.addi_course_question_value,
            style="input.TRadiobutton"
        )
        addi_courses_deny.grid(column=3, row=13, columnspan=1, sticky="e", pady=(20, 0))  # positioning

        # sets up entry box for Quantity Additional Courses Section
        self.class_quantity_box = ttk.Entry(self.scrollable_frame, textvariable=self.class_quantity)

        # Displays "Outstanding Requirements:" title to file and screen
        lbl_gpa_req = ttk.Label(self.scrollable_frame, text="Outstanding Requirements:", style="BlackSmall.TLabel")
        lbl_gpa_req.grid(column=0, row=15, columnspan=2, sticky="w", pady=(30, 5))  # text position

        # Displays outstanding core, elective, and overall gpa requirements information lines on screen
        self.lbl_out_req_sect = ttk.Label(self.scrollable_frame, text=self.out_req_sect, style="Black_txt.TLabel")
        self.lbl_out_req_sect.grid(column=0, row=16, columnspan=7, sticky="w", pady=(0, 20))  # text positioning

        # Displays instructions to select disposition of pre-reqs from degree plan
        pre_req_prompt = "Select Disposition for each Leveling Course(s) & Prerequisite(s):"
        lbl_pre_req = ttk.Label(self.scrollable_frame, text=pre_req_prompt, style="instruct_txt.TLabel")
        lbl_pre_req.grid(column=0, row=22, columnspan=6, sticky="ew", pady=(30, 0))  # text positioning

        # Displays instructions to select None if there's no prerequisite or leveling course(s)
        none_choice_instruct = "Select None if there is no Prerequisite or Leveling course(s) from the Admission Letter"
        lbl_none_choice_instruct = ttk.Label(self.scrollable_frame,
                                             text=none_choice_instruct,
                                             style="gray_subtext.TLabel")
        lbl_none_choice_instruct.grid(column=0, row=23, columnspan=7, sticky="w", pady=(5, 5))  # positioning

        # gets leveling courses to display in gui
        for lvl_class in track_class.leveling_courses:
            gui_course_list_value = lvl_class.gui_to_string()
            self.disposition_dict["dp_pre_req_class"].append(gui_course_list_value)

        # User prompted to select disposition of pre-reqs from degree plan via the GUI's drop down menu
        row_count = 23  # holds row of menu & text
        for pre_req_class in self.disposition_dict["dp_pre_req_class"]:
            row_count += 1  # holds what row widget should be
            disp_select = tk.StringVar()  # holds user disposition selection
            disp_select.set(self.disposition_dict["disp_options"][0])  # sets default value of disposition
            disp_select.trace("w", self.disposition_update)  # updates saved disposition when user changes choice
            disp_comment = tk.StringVar()  # Holds user's text entry in text box
            disp_comment.set("")  # sets default value of text entry box
            # Displays classes on screen
            lbl_disp_class = ttk.Label(self.scrollable_frame, text=pre_req_class, style="Black_txt.TLabel")
            lbl_disp_class.grid(column=0, row=row_count, columnspan=1, sticky="w", pady=(10, 0))  # positioning
            # Displays drop down menu's on screen
            option_menu = ttk.OptionMenu(
                self.scrollable_frame,
                disp_select,
                self.disposition_dict["disp_options"][0],
                *self.disposition_dict["disp_options"])
            option_menu.grid(column=1, row=row_count, columnspan=3, sticky="w", pady=(10, 0))  # positioning
            # Display Comment label for text entry box and design
            lbl_disp_comment = ttk.Label(self.scrollable_frame, text="Comment:", style="black_subtext.TLabel")
            lbl_disp_comment.grid(column=3, row=row_count, columnspan=1, sticky="e", pady=(10, 0))  # positioning
            # Displays text entry box for user to enter short comment or semester of completion
            text_entry_box = ttk.Entry(self.scrollable_frame, textvariable=disp_comment)
            text_entry_box.grid(column=4, row=row_count, columnspan=1, sticky="w", pady=(10, 0))  # positioning
            # Displays button to submit contents of text entry box
            submit_btn = ttk.Button(self.scrollable_frame, text="Submit", command=self.comment_tracker)
            submit_btn.grid(column=5, row=row_count, columnspan=1, sticky="e", pady=(10, 0))  # positioning
            self.disposition_dict["disp_selections"].append(disp_select)  # appends needed amount of StringVar()
            self.disposition_dict["opt_menu"].append(option_menu)  # Adds disposition menu with unique variable to list
            self.disposition_dict["user_course_comment"].append(disp_comment)  # adds entry StringVar() to list
            self.disposition_dict["entry_box"].append(text_entry_box)  # adds entry widget to list
            self.disposition_dict["lbl_course_widget"].append(lbl_disp_class)  # appends course name label widget
            self.disposition_dict["lbl_comment_widget"].append(lbl_disp_comment)  # appends comment label widget
            self.disposition_dict["btn_submit"].append(submit_btn)  # appends submit button widget

        final_row = row_count+1  # Holds last row value

        # Previous Page button and design
        self.previous_btn = ttk.Button(
            self.scrollable_frame,
            text="<< Previous",
            command=lambda: self.controller.show_frame("DegreePlanPage")
        )
        self.previous_btn.grid(column=0, row=final_row, columnspan=1, sticky="sw", padx=(5, 0), pady=(20, 20))

        # Saves printable audit report in File Explorer
        self.audit_report_btn = ttk.Button(
            self.scrollable_frame,
            text="Save Audit Report",
            command=lambda: self.save_file()
        )
        self.audit_report_btn.grid(column=3, row=final_row, columnspan=1, sticky="sw", padx=(0, 5), pady=20)

        # Go to homepage button and design
        self.homepage_btn = ttk.Button(
            self.scrollable_frame,
            text="Go to Homepage",
            command=lambda: self.controller.show_frame("HomepageStart")
        )
        self.homepage_btn.grid(column=4, row=final_row, columnspan=1, sticky="se", pady=(20, 20))  # positioning

    # Assigns track based on which track was selected
    def check_track(self):
        selected_track = self.student_track
        track_class = DataScience()  # default track selection
        if selected_track == "Data Sciences":
            track_class = DataScience()
        elif selected_track == "Systems":
            track_class = Systems()
        elif selected_track == "Interactive Computing":
            track_class = InteractiveComputing()
        elif selected_track == "Cyber Security":
            track_class = CyberSecurity()
        elif selected_track == "Intelligent Systems":
            track_class = IntelligentSystems()
        elif selected_track == "Networks and Telecommunications":
            track_class = Network()
        elif selected_track == "Traditional Computer Science":
            track_class = Traditional()
        elif selected_track == "Software Engineering":
            track_class = SoftwareEngineering()
        return track_class

    # Writes information to audit report file
    def generate_report_content(self, audit_report_filepath):
        document = docx.Document()  # Sets up word document
        transcript_filepath = self.transcript_path  # holds filepath of uploaded file
        transcript = Transcript(transcript_filepath)  # used to get info from transcript

        # Holds student's personal information
        self.student_name = transcript.get_name()
        self.student_id = transcript.get_id()
        self.student_plan = "Master"
        self.student_major = transcript.get_major()
        selected_track = self.student_track

        track_class = self.check_track()  # holds class for corresponding selected track
        audit_report = AuditReport(transcript_filepath, track_class)  # used to call audit report related methods

        self.core_courses = audit_report.get_core_section()  # Holds core courses displayed in audit report file
        self.elective_courses = audit_report.get_electives_section()  # Holds electives displayed in audit report file

        # sets default font
        report_style = document.styles["Normal"]
        fonts = report_style.font
        fonts.name = "Calibri"

        # Handles word doc margin
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Displays "Audit Report" title sent in file
        ar_title = document.add_paragraph()
        ar_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # centers text
        ar_title.paragraph_format.space_after = Pt(12)  # adds space after paragraph
        format_ar_title = ar_title.add_run("Audit Report")  # adds text
        format_ar_title.bold = True  # sets text to bold
        format_ar_title.font.size = Pt(14)  # sets font size

        # Displays student name in audit report word doc
        ar_student_name = document.add_paragraph()
        format_ar_student_lbl = ar_student_name.add_run("Name: ")  # adds text
        format_ar_student_lbl.bold = True  # sets text to bold
        format_ar_student_lbl.font.size = Pt(12)  # sets font size
        format_ar_student_name = ar_student_name.add_run(self.student_name)
        format_ar_student_name.font.size = Pt(12)  # sets font size

        # Displays id in audit report word doc
        format_ar_student_id_lbl = ar_student_name.add_run("               ID: ")  # adds text
        format_ar_student_id_lbl.bold = True  # sets text to bold
        format_ar_student_id_lbl.font.size = Pt(12)  # sets font size
        format_ar_student_id = ar_student_name.add_run(self.student_id)
        format_ar_student_id.font.size = Pt(12)  # sets font size
        ar_student_name.paragraph_format.space_after = Pt(0)  # removes space after paragraph

        # Displays Plan in audit report word doc
        ar_student_plan = document.add_paragraph()
        format_ar_student_plan_lbl = ar_student_plan.add_run("Plan: ")  # adds text
        format_ar_student_plan_lbl.bold = True  # sets text to bold
        format_ar_student_plan_lbl.font.size = Pt(12)  # sets font size
        format_ar_student_plan = ar_student_plan.add_run(self.student_plan)
        format_ar_student_plan.font.size = Pt(12)  # sets font size

        # Displays Student major in word doc
        format_ar_student_major_lbl = ar_student_plan.add_run("                              Major: ")  # adds text
        format_ar_student_major_lbl.bold = True  # sets text to bold
        format_ar_student_major_lbl.font.size = Pt(12)  # sets font size
        format_ar_student_major = ar_student_plan.add_run(self.student_major)
        format_ar_student_major.font.size = Pt(12)  # sets font size
        ar_student_plan.paragraph_format.space_after = Pt(0)  # removes space after paragraph

        # Displays student track in word doc
        ar_student_track = document.add_paragraph()
        format_ar_student_track_lbl = ar_student_track.add_run("Track: ")  # adds text
        format_ar_student_track_lbl.bold = True  # sets text to bold
        format_ar_student_track_lbl.font.size = Pt(12)  # sets font size
        format_ar_student_track = ar_student_track.add_run(selected_track)
        format_ar_student_track.font.size = Pt(12)  # sets font size
        ar_student_track.paragraph_format.space_after = Pt(12)  # adds space after paragraph

        # Displays "Core GPA" row in word doc
        ar_core_gpa = document.add_paragraph()
        format_ar_core_gpa_lbl = ar_core_gpa.add_run("Core GPA: ")  # adds text
        format_ar_core_gpa_lbl.bold = True  # sets text to bold
        format_ar_core_gpa_lbl.font.size = Pt(12)  # sets font size
        format_ar_core_gpa = ar_core_gpa.add_run(self.core_gpa)
        format_ar_core_gpa.font.size = Pt(12)  # sets font size
        ar_core_gpa.paragraph_format.space_after = Pt(0)  # removes space after paragraph

        # Displays "Elective GPA" row in file
        ar_elective_gpa = document.add_paragraph()
        format_ar_elective_gpa_lbl = ar_elective_gpa.add_run("Elective GPA: ")  # adds text
        format_ar_elective_gpa_lbl.bold = True  # sets text to bold
        format_ar_elective_gpa_lbl.font.size = Pt(12)  # sets font size
        format_ar_elective_gpa = ar_elective_gpa.add_run(self.elective_gpa)
        format_ar_elective_gpa.font.size = Pt(12)  # sets font size
        ar_elective_gpa.paragraph_format.space_after = Pt(0)  # removes space after paragraph

        # Displays "Overall GPA" row in file
        ar_overall_gpa = document.add_paragraph()
        format_ar_overall_gpa_lbl = ar_overall_gpa.add_run("Combined GPA: ")  # adds text
        format_ar_overall_gpa_lbl.bold = True  # sets text to bold
        format_ar_overall_gpa_lbl.font.size = Pt(12)  # sets font size
        format_ar_overall_gpa = ar_overall_gpa.add_run(self.overall_gpa)
        format_ar_overall_gpa.font.size = Pt(12)  # sets font size
        ar_overall_gpa.paragraph_format.space_after = Pt(12)  # adds space after paragraph

        # Displays "Core Courses" row in file
        ar_core_courses = document.add_paragraph()
        format_ar_core_courses_lbl = ar_core_courses.add_run("Core Courses: ")  # adds text
        format_ar_core_courses_lbl.bold = True  # sets text to bold
        format_ar_core_courses_lbl.font.size = Pt(12)  # sets font size
        format_ar_core_courses = ar_core_courses.add_run(self.core_courses)
        format_ar_core_courses.font.size = Pt(12)  # sets font size
        ar_core_courses.paragraph_format.space_after = Pt(0)  # removes space after paragraph

        # Displays "Elective Courses" row in file
        ar_elective_courses = document.add_paragraph()
        format_ar_elective_courses_lbl = ar_elective_courses.add_run("Elective Courses: ")  # adds text
        format_ar_elective_courses_lbl.bold = True  # sets text to bold
        format_ar_elective_courses_lbl.font.size = Pt(12)  # sets font size
        format_ar_elective_courses = ar_elective_courses.add_run(self.elective_courses)
        format_ar_elective_courses.font.size = Pt(12)  # sets font size
        ar_elective_courses.paragraph_format.space_after = Pt(12)  # adds space after paragraph

        # Displays "Leveling Courses and Pre-requisites from Admission Letter:" title in word doc
        pre_req_section_title = "Leveling Courses and Pre-requisites from Admission Letter:"
        ar_pre_req_title = document.add_paragraph()
        format_ar_pre_req_title_lbl = ar_pre_req_title.add_run(pre_req_section_title)  # adds text
        format_ar_pre_req_title_lbl.bold = True  # sets text to bold
        format_ar_pre_req_title_lbl.font.size = Pt(12)  # sets font size
        ar_pre_req_title.paragraph_format.space_after = Pt(12)  # adds space after paragraph

        # Displays Course and disposition of pre-reqs in file
        ar_pre_req_courses = document.add_paragraph()
        ar_pre_req_courses.paragraph_format.space_after = Pt(0)  # removes space after paragraph

        # Displays "Outstanding Requirements:" title in file
        outstanding_req_title = "Outstanding Requirements:"
        ar_out_req_title = document.add_paragraph()
        format_ar_out_req_title_lbl = ar_out_req_title.add_run(outstanding_req_title)  # adds text
        format_ar_out_req_title_lbl.bold = True  # sets text to bold
        format_ar_out_req_title_lbl.font.size = Pt(12)  # sets font size
        ar_out_req_title.paragraph_format.space_after = Pt(12)  # adds space after paragraph

        # Displays outstanding requirements section in file
        ar_out_req = document.add_paragraph()
        format_ar_out_req_gpa = ar_out_req.add_run(self.out_req_sect)  # adds text
        format_ar_out_req_gpa.font.size = Pt(12)  # sets font size

        # Displays courses and disposition and checks if the disposition option is "None"
        lvl_sect = ""  # Holds text to be printed in audit report
        course_index = 0
        # Displays contents to word doc based on disposition selection
        for lvl_course in track_class.leveling_courses:
            check_opt = self.disposition_dict["disp_selections"][course_index]
            if check_opt.get() == "Completed":
                course_id = self.disposition_dict["dp_pre_req_class"][course_index]
                courses = get_courses(transcript_filepath)
                level_course = LevelingCourse(course_id, status=LevelingCourseStatus.COMPLETED)
                print_course = level_course.to_string(courses)
                get_comment = self.disposition_dict["user_course_comment"][course_index]
                lvl_sect += print_course + " " + get_comment.get() + "\n"
            elif check_opt.get() == "Waived":
                get_comment = self.disposition_dict["user_course_comment"][course_index]
                print_lvl = self.disposition_dict["dp_pre_req_class"][course_index] + " Waived " + get_comment.get()
                lvl_sect += print_lvl + "\n"
            elif check_opt.get() == "Other":
                get_comment = self.disposition_dict["user_course_comment"][course_index]
                print_lvl = self.disposition_dict["dp_pre_req_class"][course_index] + " Other " + get_comment.get()
                lvl_sect += print_lvl + "\n"
            elif check_opt.get() == "Not required by plan or elective":
                status = " Not required by plan or elective "
                get_comment = self.disposition_dict["user_course_comment"][course_index]
                print_lvl = self.disposition_dict["dp_pre_req_class"][course_index] + status + get_comment.get()
                lvl_sect += print_lvl + "\n"
            else:
                status = " Leveling Course Marked as None "
                get_comment = self.disposition_dict["user_course_comment"][course_index]
                print_lvl = self.disposition_dict["dp_pre_req_class"][course_index] + status + get_comment.get()
                lvl_sect += print_lvl + "\n"
            course_index += 1

        format_ar_pre_req_courses = ar_pre_req_courses.add_run(lvl_sect)  # adds text
        format_ar_pre_req_courses.font.size = Pt(12)  # sets font size

        document.save(audit_report_filepath)  # saves the Word doc
